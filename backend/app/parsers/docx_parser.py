"""DOCX parser — extracts paragraphs and table text."""
from .base import BaseParser, ParsedContent
from ..config import get_settings


class DocxParser(BaseParser):
    supported_extensions = [".docx"]

    def parse(self, filepath: str) -> ParsedContent:
        result = self._make_result(filepath, "docx")
        result.extraction_method = "python-docx"
        max_chars = get_settings().max_chars_per_file
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            combined = "\n".join(parts)
            if len(combined) > max_chars:
                result.text_content = combined[:max_chars] + f"\n...[truncated, {len(combined)} chars total]"
            else:
                result.text_content = combined
        except Exception as exc:
            result.error = str(exc)
            result.text_content = f"[Could not read .docx: {exc}]"
        return result
